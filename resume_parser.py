import re
import io
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from pathlib import Path

import PyPDF2
from docx import Document
from PIL import Image

try:
    import cv2
    import numpy as np
except ImportError:
    cv2 = None
    np = None

try:
    import pytesseract
except ImportError:
    pytesseract = None


@dataclass
class ParsedResume:
    """解析后的简历数据结构"""
    name: str = ""
    email: str = ""
    phone: str = ""
    education: List[Dict[str, Any]] = field(default_factory=list)
    work_experience: List[Dict[str, Any]] = field(default_factory=list)
    skills: List[str] = field(default_factory=list)
    projects: List[Dict[str, Any]] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    raw_text: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "email": self.email,
            "phone": self.phone,
            "education": self.education,
            "work_experience": self.work_experience,
            "skills": self.skills,
            "projects": self.projects,
            "certifications": self.certifications,
        }


class ResumeParser:
    """简历解析器 - 支持PDF、DOCX、TXT格式"""
    
    def __init__(self):
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'(?:(?:\+?86)?[-\s]?)?1[3-9]\d{9}|(?:\d{3}-)?\d{8}|(?:\d{4}-)?\d{7,8}')
        self.education_keywords = ['本科', '硕士', '博士', '大专', '高中', '学士', '研究生', '博士', 'MBA', 'EMBA']
        self.skill_keywords = [
            'python', 'java', 'javascript', 'js', 'typescript', 'ts', 'go', 'golang', 'rust', 'c++', 'c#',
            'react', 'vue', 'angular', 'node', 'nodejs', 'django', 'flask', 'spring',
            'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'kafka',
            'docker', 'kubernetes', 'k8s', 'aws', 'azure', 'gcp', '阿里云',
            '机器学习', '深度学习', '人工智能', 'ai', 'nlp', '计算机视觉',
            '数据分析', '数据挖掘', '大数据', 'hadoop', 'spark', 'flink',
            '产品经理', '项目管理', '敏捷开发', 'scrum', 'devops'
        ]
    
    def parse(self, file_content: bytes, file_extension: str) -> ParsedResume:
        """根据文件类型解析简历"""
        if file_extension.lower() == '.pdf':
            raw_text = self._extract_pdf(file_content)
        elif file_extension.lower() in ['.docx', '.doc']:
            raw_text = self._extract_docx(file_content)
        elif file_extension.lower() == '.txt':
            raw_text = file_content.decode('utf-8', errors='ignore')
        elif file_extension.lower() in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            raw_text = self._extract_image(file_content)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
        
        return self._parse_text(raw_text)
    
    def _extract_pdf(self, content: bytes) -> str:
        """从PDF提取文本"""
        text = ""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PDF解析错误: {e}")
        return text
    
    def _extract_docx(self, content: bytes) -> str:
        """从DOCX提取文本"""
        text = ""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            print(f"DOCX解析错误: {e}")
        return text
    
    def _extract_image(self, content: bytes) -> str:
        """从图片提取文本（OCR）"""
        text = ""
        if not pytesseract or not cv2 or not np:
            print("OCR功能未启用，请安装tesseract-ocr, opencv-python和numpy")
            return text
        
        try:
            # 读取图片
            img = Image.open(io.BytesIO(content))
            
            # 预处理图片
            img_array = np.array(img)
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # 二值化
            _, binary = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)
            
            # 保存处理后的图片用于OCR
            processed_img = Image.fromarray(binary)
            
            # 执行OCR
            text = pytesseract.image_to_string(processed_img, lang='chi_sim+eng')
        except Exception as e:
            print(f"图片OCR错误: {e}")
        return text
    
    def _parse_text(self, text: str) -> ParsedResume:
        """解析文本内容"""
        resume = ParsedResume(raw_text=text)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # 提取基本信息
        resume.email = self._extract_email(text)
        resume.phone = self._extract_phone(text)
        resume.name = self._extract_name(lines)
        
        # 提取教育背景
        resume.education = self._extract_education(text, lines)
        
        # 提取技能
        resume.skills = self._extract_skills(text)
        
        # 提取工作经历
        resume.work_experience = self._extract_work_experience(text, lines)
        
        # 提取项目经验
        resume.projects = self._extract_projects(text, lines)
        
        # 提取证书
        resume.certifications = self._extract_certifications(text, lines)
        
        return resume
    
    def _extract_email(self, text: str) -> str:
        """提取邮箱"""
        match = self.email_pattern.search(text)
        return match.group(0) if match else ""
    
    def _extract_phone(self, text: str) -> str:
        """提取电话"""
        match = self.phone_pattern.search(text)
        return match.group(0) if match else ""
    
    def _extract_name(self, lines: List[str]) -> str:
        """提取姓名（通常在第一行或前几行）"""
        if not lines:
            return ""
        
        # 尝试从第一行提取姓名
        first_line = lines[0]
        # 去除常见标题词
        name = first_line.replace('简历', '').replace('个人简历', '').strip()
        
        # 如果第一行太长或包含特殊字符，尝试第二行
        if len(name) > 10 or any(char in name for char in ['/', '\\', '|', '-']):
            if len(lines) > 1:
                name = lines[1].strip()
        
        # 限制姓名长度
        if len(name) > 20:
            name = name[:20]
        
        return name
    
    def _extract_education(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """提取教育背景"""
        education = []
        
        # 查找教育相关部分
        education_section = self._extract_section(text, ['教育背景', '教育经历', '学历', 'Education'])
        
        if education_section:
            # 提取学校名称（常见模式）
            school_pattern = re.compile(r'(大学|学院|学校|University|College|Institute)[\u4e00-\u9fa5a-zA-Z]*')
            schools = school_pattern.findall(education_section)
            
            # 提取学历
            for keyword in self.education_keywords:
                if keyword in education_section:
                    edu_entry = {
                        "degree": keyword,
                        "school": schools[0] if schools else "",
                        "major": self._extract_major(education_section),
                        "period": self._extract_period(education_section)
                    }
                    education.append(edu_entry)
                    break
        
        return education
    
    def _extract_major(self, text: str) -> str:
        """提取专业"""
        major_pattern = re.compile(r'(专业|Major)[：:\s]*([\u4e00-\u9fa5a-zA-Z]+)')
        match = major_pattern.search(text)
        return match.group(2) if match else ""
    
    def _extract_period(self, text: str) -> str:
        """提取时间段"""
        period_pattern = re.compile(r'(\d{4}[\.\-/年]\d{1,2}[\s]*[-~至到][\s]*\d{4}[\.\-/年]\d{1,2}|\d{4}[\.\-/年][\s]*[-~至到][\s]*至今|\d{4}[\.\-/年][\s]*[-~至到][\s]*Present)')
        match = period_pattern.search(text)
        return match.group(0) if match else ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """提取技能"""
        skills = []
        text_lower = text.lower()
        
        # 查找技能部分
        skills_section = self._extract_section(text, ['技能', '专业技能', '技术栈', 'Skills', 'Technical Skills'])
        
        if skills_section:
            # 从技能部分提取关键词
            for skill in self.skill_keywords:
                if skill.lower() in skills_section.lower():
                    skills.append(skill)
        
        # 如果没有找到技能部分，在整个文本中查找
        if not skills:
            for skill in self.skill_keywords:
                if skill.lower() in text_lower:
                    skills.append(skill)
        
        return list(set(skills))  # 去重
    
    def _extract_work_experience(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """提取工作经历"""
        experiences = []
        
        work_section = self._extract_section(text, ['工作经历', '工作经验', 'Work Experience', 'Experience'])
        
        if work_section:
            # 尝试提取公司名称和职位
            company_pattern = re.compile(r'([\u4e00-\u9fa5a-zA-Z]+(?:公司|集团|科技|网络|Corp|Inc|Ltd|Company))')
            companies = company_pattern.findall(work_section)
            
            position_pattern = re.compile(r'(工程师|开发|经理|主管|总监|Engineer|Developer|Manager|Director)')
            positions = position_pattern.findall(work_section)
            
            for i, company in enumerate(companies[:3]):  # 最多取3个
                exp = {
                    "company": company,
                    "position": positions[i] if i < len(positions) else "",
                    "period": self._extract_period(work_section),
                    "description": ""
                }
                experiences.append(exp)
        
        return experiences
    
    def _extract_projects(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """提取项目经验"""
        projects = []
        
        project_section = self._extract_section(text, ['项目经验', '项目经历', 'Projects', 'Project Experience'])
        
        if project_section:
            # 简单的项目提取逻辑
            project_names = re.findall(r'[\u4e00-\u9fa5]{2,20}(?:项目|系统|平台)', project_section)
            
            for name in project_names[:3]:
                project = {
                    "name": name,
                    "role": "",
                    "period": self._extract_period(project_section),
                    "description": ""
                }
                projects.append(project)
        
        return projects
    
    def _extract_certifications(self, text: str, lines: List[str]) -> List[str]:
        """提取证书"""
        certs = []
        
        cert_section = self._extract_section(text, ['证书', '认证', 'Certifications', 'Certificates'])
        
        if cert_section:
            cert_keywords = ['PMP', 'AWS', 'Azure', 'CKA', '软考', '思科', 'Cisco', 'Oracle', '微软']
            for keyword in cert_keywords:
                if keyword in cert_section:
                    certs.append(keyword)
        
        return certs
    
    def _extract_section(self, text: str, section_names: List[str]) -> str:
        """提取特定部分的内容"""
        lines = text.split('\n')
        section_content = []
        in_section = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # 检查是否是目标部分的开始
            if any(name in line_stripped for name in section_names):
                in_section = True
                continue
            
            # 检查是否是下一个部分的开始（常见标题词）
            if in_section:
                next_section_indicators = ['教育', '工作', '项目', '技能', '证书', '自我评价', '获奖', '语言']
                if any(indicator in line_stripped for indicator in next_section_indicators) and len(line_stripped) < 20:
                    if not any(name in line_stripped for name in section_names):
                        break
                section_content.append(line)
        
        return '\n'.join(section_content)
